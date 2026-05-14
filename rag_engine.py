import os
import io
import re
import logging

from dotenv import load_dotenv
import google.generativeai as genai
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

load_dotenv()
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
logger = logging.getLogger(__name__)

_chunks = []
_vectorizer = None
_matrix = None


# ═══════════════════════════════════════════════════════
# TEXT EXTRACTION
# ═══════════════════════════════════════════════════════

def extract_text(file_bytes, filename):
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
    if ext == "pdf":
        return _extract_pdf(file_bytes, filename)
    if ext == "docx":
        return _extract_docx(file_bytes)
    if ext == "csv":
        return _extract_csv(file_bytes)
    return _extract_plaintext(file_bytes, filename)


def _extract_pdf(file_bytes, filename):
    text = ""

    # Method 1: PyMuPDF (best — works on most PDFs including many scanned ones)
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = []
        for page in doc:
            t = page.get_text("text")
            if t and t.strip():
                pages.append(t.strip())
        doc.close()
        text = "\n".join(pages).strip()
        if text and len(text) > 50:
            logger.info("PyMuPDF extracted %d chars from '%s'", len(text), filename)
            return text
    except Exception as e:
        logger.warning("PyMuPDF failed: %s", e)

    # Method 2: PyPDF2
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        pages = []
        for page in reader.pages:
            try:
                t = page.extract_text()
                if t and t.strip():
                    pages.append(t.strip())
            except Exception:
                continue
        text = "\n".join(pages).strip()
        if text and len(text) > 50:
            logger.info("PyPDF2 extracted %d chars from '%s'", len(text), filename)
            return text
    except Exception as e:
        logger.warning("PyPDF2 failed: %s", e)

    # If both fail → scanned PDF
    raise RuntimeError(
        "SCANNED_PDF"
    )


def _extract_docx(file_bytes):
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        lines = []
        for p in doc.paragraphs:
            if p.text.strip():
                lines.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    lines.append(row_text)
        text = "\n".join(lines).strip()
        if text:
            return text
        raise RuntimeError("DOCX appears to be empty.")
    except ImportError:
        raise RuntimeError("Run: pip install python-docx")


def _extract_csv(file_bytes):
    try:
        import csv
        text_io = io.StringIO(file_bytes.decode("utf-8", errors="replace"))
        reader = csv.reader(text_io)
        rows = []
        for row in reader:
            cleaned = ", ".join(cell.strip() for cell in row if cell.strip())
            if cleaned:
                rows.append(cleaned)
        return "\n".join(rows).strip()
    except Exception:
        return file_bytes.decode("utf-8", errors="replace").strip()


def _extract_plaintext(file_bytes, filename):
    for encoding in ["utf-8", "latin-1", "cp1252"]:
        try:
            text = file_bytes.decode(encoding, errors="replace").strip()
            if text:
                return text
        except Exception:
            continue
    raise RuntimeError(f"Could not read '{filename}'.")


# ═══════════════════════════════════════════════════════
# CHUNKING
# ═══════════════════════════════════════════════════════

def chunk_text(text, chunk_size=300, overlap=50):
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r' {3,}', ' ', text)
    words = text.split()
    if not words:
        return []
    chunks = []
    i = 0
    while i < len(words):
        chunk = " ".join(words[i: i + chunk_size]).strip()
        if len(chunk) > 10:
            chunks.append(chunk)
        i += chunk_size - overlap
        if i >= len(words):
            break
    return chunks


# ═══════════════════════════════════════════════════════
# VECTOR STORE
# ═══════════════════════════════════════════════════════

def store_document(filename, text):
    global _chunks, _vectorizer, _matrix
    _chunks = [c for c in _chunks if c["filename"] != filename]
    new_chunks = chunk_text(text)
    if not new_chunks:
        raise RuntimeError("Could not create any chunks from this file.")
    for chunk in new_chunks:
        _chunks.append({"filename": filename, "text": chunk})
    _rebuild_index()
    logger.info("Stored '%s': %d chunks. Total: %d", filename, len(new_chunks), len(_chunks))
    return len(new_chunks)


def _rebuild_index():
    global _vectorizer, _matrix, _chunks
    if not _chunks:
        _vectorizer = None
        _matrix = None
        return
    all_texts = [c["text"] for c in _chunks]
    _vectorizer = TfidfVectorizer(
        ngram_range=(1, 2),
        max_features=20000,
        sublinear_tf=True
    )
    _matrix = _vectorizer.fit_transform(all_texts)
    logger.info("Index rebuilt: %d chunks", len(_chunks))


def get_store_stats():
    unique_files = list({c["filename"] for c in _chunks})
    return {"total_chunks": len(_chunks), "documents": unique_files}


def delete_document(filename):
    """Delete one specific document from the store."""
    global _chunks
    _chunks = [c for c in _chunks if c["filename"] != filename]
    _rebuild_index()
    logger.info("Deleted '%s'. Remaining chunks: %d", filename, len(_chunks))


def clear_store():
    global _chunks, _vectorizer, _matrix
    _chunks = []
    _vectorizer = None
    _matrix = None
    logger.info("Store cleared")


# ═══════════════════════════════════════════════════════
# RETRIEVAL — with optional file filter
# ═══════════════════════════════════════════════════════

def retrieve_context(question, top_k=6, filename_filter=None):
    """
    filename_filter: if provided, only search chunks from that file.
    If None, search across ALL uploaded files.
    """
    global _vectorizer, _matrix, _chunks

    if not _chunks:
        logger.warning("No chunks in store")
        return []

    if _vectorizer is None or _matrix is None:
        _rebuild_index()
        if _vectorizer is None:
            return []

    # Apply file filter if specified
    if filename_filter:
        filtered = [c for c in _chunks if c["filename"] == filename_filter]
        if not filtered:
            logger.warning("No chunks found for file: %s", filename_filter)
            return []
        # Build a temporary index just for this file
        texts = [c["text"] for c in filtered]
        temp_vectorizer = TfidfVectorizer(
            ngram_range=(1, 2),
            max_features=20000,
            sublinear_tf=True
        )
        temp_matrix = temp_vectorizer.fit_transform(texts)
        q_vec = temp_vectorizer.transform([question])
        scores = cosine_similarity(q_vec, temp_matrix).flatten()
        top_indices = scores.argsort()[::-1][:top_k]
        return [{"filename": filtered[idx]["filename"],
                 "text": filtered[idx]["text"],
                 "score": float(scores[idx])} for idx in top_indices]

    # Search across all files
    try:
        q_vec = _vectorizer.transform([question])
        scores = cosine_similarity(q_vec, _matrix).flatten()
        top_indices = scores.argsort()[::-1][:top_k]
        results = []
        for idx in top_indices:
            results.append({
                **_chunks[idx],
                "score": float(scores[idx])
            })
        logger.info("Retrieved %d chunks (max score: %.4f)", len(results),
                    scores.max() if len(scores) > 0 else 0)
        return results
    except Exception as e:
        logger.error("Retrieval error: %s", e)
        return _chunks[:top_k]


# ═══════════════════════════════════════════════════════
# ANSWER GENERATION
# ═══════════════════════════════════════════════════════

def generate_answer(question, context_chunks, conversation_history=None, filename_filter=None):
    try:
        if context_chunks:
            context_text = "\n\n---\n\n".join(
                f'[Part {i+1} from "{c["filename"]}"]\n{c["text"]}'
                for i, c in enumerate(context_chunks)
            )

            # Tell AI which file to focus on if filter is set
            file_instruction = ""
            if filename_filter:
                file_instruction = (
                    f"The user wants answers specifically from the file: '{filename_filter}'.\n"
                    "Only use information from that file.\n\n"
                )

            prompt = (
                "You are a helpful AI assistant that answers questions "
                "based on uploaded documents.\n\n"
                f"{file_instruction}"
                "RULES:\n"
                "1. Answer using the document context below.\n"
                "2. Give detailed, helpful answers — summaries, short notes, "
                "key points, explanations as needed.\n"
                "3. If asked for short notes or summary, give well-structured "
                "bullet points with headings.\n"
                "4. Use markdown formatting (bullets, bold, headings).\n"
                "5. If the answer is not in the context, say so clearly.\n\n"
                f"DOCUMENT CONTEXT:\n{context_text}\n\n"
                f"QUESTION: {question}"
            )
        else:
            stats = get_store_stats()
            if stats["total_chunks"] > 0:
                # Files ARE uploaded but retrieval returned nothing
                prompt = (
                    "You are a helpful AI assistant.\n"
                    f"The user has uploaded these files: {stats['documents']}\n"
                    "But the specific question could not be matched to any content.\n"
                    "Tell the user you have their files but couldn't find relevant "
                    "content for this specific question. Ask them to rephrase.\n\n"
                    f"User question: {question}"
                )
            else:
                prompt = (
                    "You are a helpful AI assistant.\n"
                    "No document has been uploaded yet.\n"
                    "Ask the user to upload a file using the upload button.\n\n"
                    f"User message: {question}"
                )

        models_to_try = [
            "models/gemini-2.5-flash",
            "models/gemini-2.0-flash",
            "models/gemini-2.0-flash-lite",
            "models/gemini-flash-latest",
        ]

        last_error = None
        for model_name in models_to_try:
            try:
                model = genai.GenerativeModel(model_name)
                response = model.generate_content(prompt)
                return response.text
            except Exception as e:
                last_error = e
                logger.warning("Model %s failed: %s", model_name, e)
                continue

        return f"Error: {str(last_error)}"

    except Exception as e:
        logger.exception("generate_answer error")
        return f"Error generating response: {str(e)}"